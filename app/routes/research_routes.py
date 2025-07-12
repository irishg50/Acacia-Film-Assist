from flask import Blueprint, request, jsonify, current_app, send_file
from flask_login import login_required, current_user
from app.services.perplexity_service import PerplexityService
from app.services.auth_decorators import login_required
from app.extensions import db
from app.models.models import ResearchSession
from datetime import datetime
import json
from io import BytesIO
from docx import Document

research_bp = Blueprint('research_bp', __name__)

@research_bp.route('/api/research', methods=['POST'])
@login_required
def research_topic():
    """Handle research requests using Perplexity API"""
    try:
        data = request.get_json()
        topic = data.get('topic')
        focus_areas = data.get('focus_areas', [])
        project_id = data.get('project_id')

        if not topic:
            return jsonify({"error": "Topic is required"}), 400

        if not project_id:
            return jsonify({"error": "Project ID is required"}), 400

        # Initialize Perplexity service
        perplexity = PerplexityService()
        
        # Perform research
        research_data = perplexity.research_topic(topic, focus_areas)
        
        # Save research session to database
        research_session = ResearchSession(
            user_id=current_user.id,
            project_id=project_id,
            topic=topic,
            focus_areas=str(focus_areas),
            research_content=research_data['content'],
            created_at=datetime.utcnow(),
            perplexity_response=json.dumps(research_data)
        )
        db.session.add(research_session)
        db.session.commit()

        return jsonify({
            "status": "success",
            "research": research_data,
            "session_id": research_session.id
        })

    except Exception as e:
        current_app.logger.error(f"Research error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@research_bp.route('/api/research/history', methods=['GET'])
@login_required
def get_research_history():
    """Get research history for the current project"""
    try:
        project_id = request.args.get('project_id', type=int)
        if not project_id:
            return jsonify({"error": "Project ID is required"}), 400

        research_sessions = ResearchSession.query.filter_by(
            user_id=current_user.id,
            project_id=project_id
        ).order_by(ResearchSession.created_at.desc()).all()

        return jsonify({
            "status": "success",
            "research_history": [{
                "id": session.id,
                "topic": session.topic,
                "focus_areas": eval(session.focus_areas),
                "created_at": session.created_at.isoformat(),
                "preview": session.research_content[:200] + "..." if len(session.research_content) > 200 else session.research_content,
                "research_content": session.research_content
            } for session in research_sessions]
        })

    except Exception as e:
        current_app.logger.error(f"Error fetching research history: {str(e)}")
        return jsonify({"error": str(e)}), 500

@research_bp.route('/api/research/<int:session_id>/download', methods=['GET'])
@login_required
def download_research_report(session_id):
    session = ResearchSession.query.filter_by(id=session_id, user_id=current_user.id).first_or_404()
    if not session.perplexity_response:
        return jsonify({'error': 'No Perplexity response found for this research session.'}), 404
    try:
        data = json.loads(session.perplexity_response)
        doc = Document()
        doc.add_heading(session.topic, 0)
        doc.add_paragraph(f"Created: {session.created_at.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("")
        # Main content
        content = data.get('content') or data.get('research_content') or ''
        doc.add_paragraph(content)
        # Citations
        citations = data.get('citations')
        if citations:
            doc.add_heading('Citations', level=1)
            for c in citations:
                doc.add_paragraph(c, style='List Number')
        # Search Results
        search_results = data.get('search_results')
        if search_results:
            doc.add_heading('Search Results', level=1)
            for r in search_results:
                p = doc.add_paragraph()
                p.add_run(r.get('title', 'Untitled')).bold = True
                if r.get('url'):
                    p.add_run(f" ({r['url']})")
                if r.get('date'):
                    p.add_run(f" [{r['date']}]")
        # Save to BytesIO
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"Research_Report_{session.topic[:40].replace(' ', '_')}.docx"
        return send_file(buf, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({'error': f'Failed to generate report: {str(e)}'}), 500 